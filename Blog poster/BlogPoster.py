import requests
from bs4 import BeautifulSoup
from docx import Document

# List of blog URLs to scrape
blog_urls = [
    'https://exampleblog1.com',
    'https://exampleblog2.com',
    'https://exampleblog3.com'
]

# Function to scrape blog data and store in Word format
def scrape_and_store_data(url):
    # Make a GET request to the blog URL
    response = requests.get(url)
    
    # Parse the HTML content using BeautifulSoup
    soup = BeautifulSoup(response.content, 'html.parser')
    
    # Extract the relevant data from the HTML
    # Replace the following code with your own logic to extract the desired data
    
    # Example: Extracting the blog title
    blog_title = soup.find('h1').text
    
    # Example: Extracting the blog content
    blog_content = soup.find('div', class_='content').text
    
    # Create a new Word document
    document = Document()
    
    # Add the blog title as the document title
    document.add_heading(blog_title, level=1)
    
    # Add the blog content as the document content
    document.add_paragraph(blog_content)
    
    # Save the document as a Word file
    document.save(f'{blog_title}.docx')

# Scrape data from each blog URL and store in separate Word files
for url in blog_urls:
    scrape_and_store_data(url)